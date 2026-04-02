from __future__ import annotations

import json
import sys
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_NAME = "Arial"
ARABIC_DIGITS = str.maketrans("0123456789", "٠١٢٣٤٥٦٧٨٩")


def to_arabic_digits(value: Any) -> str:
    return str(value if value is not None else "").translate(ARABIC_DIGITS)


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return to_arabic_digits(value)


def make_element(tag: str, attribs: dict[str, str] | None = None):
    element = OxmlElement(tag)
    if attribs:
        for key, value in attribs.items():
            element.set(qn(key), value)
    return element


def set_para_rtl(paragraph):
    ppr = paragraph._p.get_or_add_pPr()

    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        ppr.insert(0, make_element("w:bidi"))

    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = make_element("w:jc", {"w:val": "right"})
        ppr.append(jc)
    else:
        jc.set(qn("w:val"), "right")


def set_run_rtl(run, font_name: str = FONT_NAME, font_size: int = 12):
    rpr = run._r.get_or_add_rPr()

    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rpr.append(make_element("w:rtl"))

    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = make_element("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_doc_rtl(doc: Document):
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        sect_pr = make_element("w:sectPr")
        body.append(sect_pr)

    if sect_pr.find(qn("w:bidi")) is None:
        sect_pr.insert(0, make_element("w:bidi"))


def configure_section(section, landscape: bool = False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)


def set_default_style_rtl(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    rpr = normal._element.rPr
    if rpr is None:
        rpr = make_element("w:rPr")
        normal._element.insert(0, rpr)

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = make_element("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    ppr = normal._element.pPr
    if ppr is None:
        ppr = make_element("w:pPr")
        normal._element.append(ppr)

    if ppr.find(qn("w:bidi")) is None:
        ppr.insert(0, make_element("w:bidi"))

    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = make_element("w:jc", {"w:val": "right"})
        ppr.append(jc)
    else:
        jc.set(qn("w:val"), "right")


def set_table_rtl(table):
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = make_element("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(make_element("w:bidiVisual"))


def set_cell_shading(cell, fill: str = "FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = make_element("w:shd", {"w:val": "clear", "w:fill": fill})
        tc_pr.append(shd)
    else:
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)


def add_text_run(paragraph, text: Any, *, bold: bool = False, size: int = 12, color: str | None = None):
    run = paragraph.add_run(normalize_text(text))
    set_run_rtl(run, font_size=size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_rtl_paragraph(
    doc: Document,
    text: Any,
    *,
    bold: bool = False,
    size: int = 12,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    space_before: int = 0,
    space_after: int = 0,
    color: str | None = None,
):
    paragraph = doc.add_paragraph()
    set_para_rtl(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    add_text_run(paragraph, text, bold=bold, size=size, color=color)
    return paragraph


def add_bottom_line_paragraph(doc: Document, text: Any = " ", *, size: int = 12, space_after: int = 0):
    paragraph = doc.add_paragraph()
    set_para_rtl(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(space_after)
    add_text_run(paragraph, text, size=size)
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = make_element("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = make_element(
            "w:bottom",
            {"w:val": "single", "w:sz": "4", "w:space": "1", "w:color": "000000"},
        )
        pbdr.append(bottom)
    return paragraph


def add_single_cell_table(doc: Document, text: Any, *, size: int = 14, bold: bool = True):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_rtl(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFFFFF")
    paragraph = cell.paragraphs[0]
    set_para_rtl(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(paragraph, text, bold=bold, size=size)
    return table


def add_metadata_table(doc: Document, exam_meta: dict[str, Any]):
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    set_table_rtl(table)

    rows = [
        [
            ("المادة", exam_meta.get("subject", "—")),
            ("الصف / الفئة", exam_meta.get("className", "—")),
            ("التاريخ", exam_meta.get("date", "—")),
            ("المدة", exam_meta.get("duration", "—")),
        ],
        [
            ("الدرجة الكلية", exam_meta.get("totalMarks", "—")),
            ("عدد الأسئلة", exam_meta.get("totalQuestions", "—")),
            ("الفصل الدراسي", exam_meta.get("term", "—")),
            ("العام الدراسي", exam_meta.get("academicYear", "—")),
        ],
    ]

    for row_index, row in enumerate(rows):
        for col_index, (label, value) in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "FFFFFF")
            paragraph = cell.paragraphs[0]
            set_para_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_after = Pt(0)
            add_text_run(paragraph, f"{label}: ", bold=True, size=11)
            add_text_run(paragraph, value, size=11)

    return table


def add_student_info_table(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_rtl(table)

    labels = [
        "اسم الطالب: ____________________",
        "الشعبة: ____________________",
    ]
    for index, text in enumerate(labels):
        cell = table.cell(0, index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "FFFFFF")
        paragraph = cell.paragraphs[0]
        set_para_rtl(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        add_text_run(paragraph, text, bold=True, size=11)

    return table


def add_section_title_table(doc: Document, title: str, marks: int | None = None):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_rtl(table)

    content_cell = table.cell(0, 0)
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(content_cell, "FFFFFF")
    content_para = content_cell.paragraphs[0]
    set_para_rtl(content_para)
    content_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text_run(content_para, normalize_text(title), bold=True, size=12)

    score_cell = table.cell(0, 1)
    score_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(score_cell, "FFFFFF")
    score_para = score_cell.paragraphs[0]
    set_para_rtl(score_para)
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(score_para, "الدرجة", bold=True, size=11)
    if marks is not None:
        score_para_2 = score_cell.add_paragraph()
        set_para_rtl(score_para_2)
        score_para_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_para_2.paragraph_format.space_before = Pt(0)
        score_para_2.paragraph_format.space_after = Pt(0)
        add_text_run(score_para_2, to_arabic_digits(marks), bold=True, size=11)

    return table


def add_question_header(doc: Document, number: int, text: str, question_type: str):
    paragraph = doc.add_paragraph()
    set_para_rtl(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(4)

    add_text_run(paragraph, f"{to_arabic_digits(number)}- ", bold=True, size=12)
    add_text_run(paragraph, text, bold=True, size=12)

    if question_type == "true_false":
        add_text_run(paragraph, " (              )", bold=True, size=12)

    return paragraph


def add_mcq_options(doc: Document, options: list[dict[str, Any]], correct_index: int | None = None):
    for index, option in enumerate(options):
        paragraph = doc.add_paragraph()
        set_para_rtl(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(2)

        is_correct = correct_index is not None and index == correct_index
        add_text_run(
            paragraph,
            f"{option.get('label', '')}) ",
            bold=is_correct,
            size=11,
            color="1A7C1A" if is_correct else None,
        )
        add_text_run(
            paragraph,
            option.get("text", ""),
            bold=is_correct,
            size=11,
            color="1A7C1A" if is_correct else None,
        )
        if is_correct:
            add_text_run(paragraph, "  ✓", bold=True, size=10, color="1A7C1A")


def add_blank_answer_lines(doc: Document, line_count: int):
    for _ in range(line_count):
        add_bottom_line_paragraph(doc, " ", size=11, space_after=6)


def add_answer_key_block(doc: Document, question: dict[str, Any]):
    question_type = question.get("type")

    if question_type == "mcq":
        add_mcq_options(
            doc,
            question.get("options", []),
            correct_index=question.get("correctIndex"),
        )
        return

    if question_type == "true_false":
        correct = "صح" if question.get("correctAnswer") is True else "خطأ"
        add_rtl_paragraph(
            doc,
            f"الإجابة الصحيحة: {correct}",
            size=11,
            color="1A7C1A",
            space_after=3,
        )
        return

    answer_text = question.get("answerText") or question.get("correctAnswerText") or "—"
    add_rtl_paragraph(
        doc,
        f"الإجابة النموذجية: {answer_text}",
        size=11,
        space_after=3,
    )
    rubric = question.get("rubric") if isinstance(question.get("rubric"), list) else []
    if question_type == "essay" and rubric:
        add_rtl_paragraph(doc, "معايير التصحيح:", bold=True, size=11, space_after=2)
        for item in rubric:
            add_rtl_paragraph(doc, f"• {item}", size=10, space_after=0)


def add_blueprint_table(doc: Document, blueprint: dict[str, Any]):
    cells = blueprint.get("cells") if isinstance(blueprint, dict) else None
    if not isinstance(cells, list) or not cells:
        return

    headers = [
        "الدرس",
        "المستوى",
        "وزن الدرس",
        "وزن المستوى",
        "وزن الخلية",
        "الأسئلة",
        "الدرجات",
    ]

    table = doc.add_table(rows=1 + len(cells), cols=len(headers))
    table.style = "Table Grid"
    set_table_rtl(table)

    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "FFFFFF")
        paragraph = cell.paragraphs[0]
        set_para_rtl(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(paragraph, header, bold=True, size=10)

    for row_index, cell_data in enumerate(cells, start=1):
        values = [
            cell_data.get("lesson_name", ""),
            cell_data.get("level_label", ""),
            cell_data.get("topic_weight", ""),
            cell_data.get("level_weight", ""),
            cell_data.get("cell_weight", ""),
            cell_data.get("question_count", ""),
            cell_data.get("cell_marks", ""),
        ]

        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            set_para_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_run(paragraph, value, size=10)


def build_paper_or_answer_key(doc: Document, view_model: dict[str, Any], *, answer_key: bool):
    sections = view_model.get("sections") or []
    if not sections:
        add_rtl_paragraph(doc, "لا توجد أسئلة.", size=12, space_after=4)
        return

    for section_index, section in enumerate(sections):
        title = f"السؤال {['الأول', 'الثاني', 'الثالث', 'الرابع', 'الخامس'][section_index] if section_index < 5 else to_arabic_digits(section_index + 1)} : {section.get('title', '')}"
        questions = section.get("questions") or []
        section_marks = 0
        for question in questions:
            try:
                section_marks += int(question.get("marks") or 0)
            except (TypeError, ValueError):
                continue

        add_section_title_table(doc, title, marks=section_marks)

        for question in questions:
            add_question_header(
                doc,
                question.get("displayNumber", question.get("number", 1)),
                question.get("text", ""),
                question.get("type", ""),
            )

            q_type = question.get("type")
            if q_type == "mcq":
                add_mcq_options(
                    doc,
                    question.get("options", []),
                    correct_index=question.get("correctIndex") if answer_key else None,
                )
            elif q_type == "true_false":
                if answer_key:
                    correct = "صح" if question.get("correctAnswer") is True else "خطأ"
                    add_rtl_paragraph(
                        doc,
                        f"الإجابة الصحيحة: {correct}",
                        size=11,
                        color="1A7C1A",
                        space_after=2,
                    )
                else:
                    add_blank_answer_lines(doc, 1)
            elif q_type in ("short_answer", "essay"):
                if answer_key:
                    add_answer_key_block(doc, question)
                else:
                    add_blank_answer_lines(doc, 5 if q_type == "essay" else 2)
            else:
                if answer_key:
                    add_answer_key_block(doc, question)
                else:
                    add_blank_answer_lines(doc, 2)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_answer_form(doc: Document, view_model: dict[str, Any]):
    sections = view_model.get("sections") or []
    objective_sections = [
        section
        for section in sections
        if section.get("id") in {"true_false", "mcq"}
    ]

    if not objective_sections:
        add_rtl_paragraph(doc, "لا توجد أسئلة موضوعية في هذا الاختبار.", size=11, space_after=4)
        return

    add_rtl_paragraph(
        doc,
        "ظلِّل خيارًا واحدًا فقط لكل سؤال، ولا تظلِّل أكثر من خيار واحد لنفس السؤال.",
        size=11,
        space_after=4,
    )

    for section in objective_sections:
        add_section_title_table(doc, section.get("title", ""), marks=None)
        questions = section.get("questions") or []
        answer_labels = ["أ", "ب", "ج", "د"] if section.get("id") == "mcq" else ["صح", "خطأ"]

        table = doc.add_table(rows=1 + len(questions), cols=1 + len(answer_labels))
        table.style = "Table Grid"
        set_table_rtl(table)

        header_cells = ["رقم السؤال", *answer_labels]
        for col_index, header in enumerate(header_cells):
            cell = table.cell(0, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "E8E8E8")
            paragraph = cell.paragraphs[0]
            set_para_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_run(paragraph, header, bold=True, size=10)

        for row_index, question in enumerate(questions, start=1):
            first_cell = table.cell(row_index, 0)
            first_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = first_cell.paragraphs[0]
            set_para_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_run(
                paragraph,
                question.get("displayNumber", question.get("number", row_index)),
                size=10,
            )

            for col_index in range(1, 1 + len(answer_labels)):
                cell = table.cell(row_index, col_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                set_para_rtl(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text_run(paragraph, " ", size=10)


def add_document_header(doc: Document, exam_meta: dict[str, Any], title: str):
    add_rtl_paragraph(doc, "الجمهورية اليمنية", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_paragraph(doc, "وزارة التربية والتعليم", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_paragraph(doc, "محافظة عدن", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    if exam_meta.get("schoolName"):
        add_rtl_paragraph(
            doc,
            f"مدرسة: {exam_meta.get('schoolName')}",
            bold=True,
            size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    add_rtl_paragraph(doc, title, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def build_document(payload: dict[str, Any]) -> Document:
    doc = Document()
    configure_section(doc.sections[0], landscape=False)
    set_default_style_rtl(doc)
    set_doc_rtl(doc)

    view_model = payload.get("viewModel") or {}
    exam_meta = view_model.get("examMeta") or {}
    export_type = payload.get("type") or "questions_only"
    blueprint = payload.get("blueprint")

    if export_type == "answer_key":
        title = "نموذج الإجابات (معلم)"
    elif export_type == "answer_form":
        title = "نموذج الإجابات"
    else:
        title = "ورقة الاختبار"

    if exam_meta.get("title"):
        if export_type == "answer_key":
            title = f"{exam_meta.get('title')} - نموذج الإجابات (معلم)"
        elif export_type == "answer_form":
            title = f"{exam_meta.get('title')} - نموذج الإجابات"
        else:
            title = exam_meta.get("title")

    add_document_header(doc, exam_meta, title)
    add_metadata_table(doc, exam_meta)

    if export_type == "answer_form":
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        add_student_info_table(doc)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        build_answer_form(doc, view_model)
    else:
        if export_type == "answer_key" and isinstance(blueprint, dict):
            add_rtl_paragraph(doc, "مصفوفة جدول المواصفات", bold=True, size=12, space_before=2, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_blueprint_table(doc, blueprint)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        build_paper_or_answer_key(doc, view_model, answer_key=export_type == "answer_key")

    return doc


def main() -> int:
    try:
        if len(sys.argv) < 2:
            raise ValueError("output path argument is required")

        output_path = sys.argv[1]
        raw_payload = sys.stdin.read()
        payload = json.loads(raw_payload) if raw_payload.strip() else {}

        doc = build_document(payload)
        doc.save(output_path)

        print(json.dumps({"success": True, "outputPath": output_path}))
        return 0
    except Exception as error:  # pragma: no cover - surfaced to python-shell
        print(json.dumps({"success": False, "error": str(error)}))
        raise


if __name__ == "__main__":
    raise SystemExit(main())
